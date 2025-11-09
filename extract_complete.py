"""
Extrator COMPLETO de template Word
Extrai: cabeçalhos, rodapés, logos, formas, imagens, cores, bordas
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import zipfile
import xml.etree.ElementTree as ET
import base64

TEMPLATE_PATH = Path('models') / 'modelo homologação.docx'


def extract_header_footer(doc_path):
    """Extrai cabeçalhos e rodapés do documento"""
    print("\n" + "="*70)
    print("📋 EXTRAINDO CABEÇALHOS E RODAPÉS")
    print("="*70 + "\n")
    
    doc = Document(doc_path)
    
    # Cabeçalhos
    for i, section in enumerate(doc.sections, 1):
        print(f"SEÇÃO {i}:")
        print("-"*70)
        
        # Cabeçalho
        if section.header:
            print(f"\n📄 CABEÇALHO:")
            for para in section.header.paragraphs:
                if para.text.strip():
                    print(f"  Texto: {para.text}")
                    print(f"  Alinhamento: {para.alignment}")
            
            # Tabelas no cabeçalho
            for j, table in enumerate(section.header.tables, 1):
                print(f"\n  📊 Tabela {j} no cabeçalho:")
                print(f"    Linhas: {len(table.rows)}, Colunas: {len(table.columns)}")
                for row_idx, row in enumerate(table.rows):
                    print(f"    Linha {row_idx}:")
                    for col_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            print(f"      [{row_idx},{col_idx}]: {cell.text[:50]}")
        
        # Rodapé
        if section.footer:
            print(f"\n📄 RODAPÉ:")
            for para in section.footer.paragraphs:
                if para.text.strip():
                    print(f"  Texto: {para.text}")
                    print(f"  Alinhamento: {para.alignment}")
            
            # Tabelas no rodapé
            for j, table in enumerate(section.footer.tables, 1):
                print(f"\n  📊 Tabela {j} no rodapé:")
                print(f"    Linhas: {len(table.rows)}, Colunas: {len(table.columns)}")
                for row_idx, row in enumerate(table.rows):
                    print(f"    Linha {row_idx}:")
                    for col_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            print(f"      [{row_idx},{col_idx}]: {cell.text[:50]}")
        
        print()


def extract_images_advanced(doc_path):
    """Extrai TODAS as imagens incluindo as em cabeçalhos/rodapés/formas"""
    print("\n" + "="*70)
    print("🖼️ EXTRAINDO TODAS AS IMAGENS")
    print("="*70 + "\n")
    
    images = []
    
    # Abrir DOCX como ZIP
    with zipfile.ZipFile(doc_path, 'r') as zip_ref:
        # Listar todos os arquivos
        for file_info in zip_ref.filelist:
            if 'media/' in file_info.filename:
                print(f"✅ Imagem encontrada: {file_info.filename}")
                
                # Ler imagem
                img_data = zip_ref.read(file_info.filename)
                
                # Converter para base64
                base64_img = base64.b64encode(img_data).decode('utf-8')
                
                # Detectar MIME
                ext = file_info.filename.split('.')[-1].lower()
                mime_types = {
                    'png': 'image/png',
                    'jpg': 'image/jpeg',
                    'jpeg': 'image/jpeg',
                    'gif': 'image/gif',
                    'bmp': 'image/bmp',
                    'emf': 'image/x-emf',
                    'wmf': 'image/x-wmf',
                }
                mime = mime_types.get(ext, 'image/png')
                
                images.append({
                    'filename': file_info.filename,
                    'name': Path(file_info.filename).name,
                    'mime': mime,
                    'base64': f"data:{mime};base64,{base64_img}",
                    'size': file_info.file_size
                })
                
                print(f"  Tipo: {mime}")
                print(f"  Tamanho: {file_info.file_size:,} bytes")
                print()
    
    return images


def extract_shapes_and_drawing(doc_path):
    """Extrai informações sobre formas e desenhos"""
    print("\n" + "="*70)
    print("🎨 EXTRAINDO FORMAS E DESENHOS")
    print("="*70 + "\n")
    
    # Abrir como ZIP e ler XML
    with zipfile.ZipFile(doc_path, 'r') as zip_ref:
        # Ler document.xml
        doc_xml = zip_ref.read('word/document.xml')
        root = ET.fromstring(doc_xml)
        
        # Namespace
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
            'v': 'urn:schemas-microsoft-com:vml',
        }
        
        # Procurar por desenhos
        drawings = root.findall('.//w:drawing', namespaces)
        print(f"📐 Total de desenhos encontrados: {len(drawings)}")
        
        for i, drawing in enumerate(drawings, 1):
            print(f"\nDesenho {i}:")
            # Tentar extrair dimensões
            extents = drawing.findall('.//wp:extent', namespaces)
            for extent in extents:
                cx = extent.get('cx', 'N/A')
                cy = extent.get('cy', 'N/A')
                print(f"  Dimensões: {cx} x {cy} EMUs")
        
        # Procurar por formas VML (Word legado)
        vml_shapes = root.findall('.//v:shape', namespaces)
        print(f"\n🔷 Total de formas VML encontradas: {len(vml_shapes)}")
        
        for i, shape in enumerate(vml_shapes, 1):
            print(f"\nForma VML {i}:")
            style = shape.get('style', 'N/A')
            print(f"  Estilo: {style[:100]}")
            shape_type = shape.get('type', 'N/A')
            print(f"  Tipo: {shape_type}")


def save_images_to_assets(images):
    """Salva imagens na pasta assets"""
    assets_dir = Path('assets') / 'extracted_images'
    assets_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"\n💾 Salvando imagens em: {assets_dir}\n")
    
    saved_images = []
    for img in images:
        filepath = assets_dir / img['name']
        
        # Decodificar base64
        import re
        base64_data = re.sub('^data:image/.+;base64,', '', img['base64'])
        image_data = base64.b64decode(base64_data)
        
        # Salvar
        with open(filepath, 'wb') as f:
            f.write(image_data)
        
        print(f"✅ {img['name']} salvo")
        saved_images.append({
            **img,
            'saved_path': str(filepath)
        })
    
    return saved_images


def generate_python_code_with_images(images):
    """Gera código Python com imagens em base64"""
    print("\n" + "="*70)
    print("📝 GERANDO CÓDIGO PYTHON")
    print("="*70 + "\n")
    
    with open('extracted_images_data.py', 'w', encoding='utf-8') as f:
        f.write('"""\n')
        f.write('Imagens extraídas do documento Word\n')
        f.write('Use estas strings base64 no HTML\n')
        f.write('"""\n\n')
        
        f.write('DOCUMENT_IMAGES = {\n')
        for i, img in enumerate(images, 1):
            var_name = img['name'].replace('.', '_').replace('-', '_')
            f.write(f'    # {img["filename"]} ({img["size"]:,} bytes)\n')
            f.write(f'    "{var_name}": "{img["base64"]}",\n\n')
        f.write('}\n\n')
        
        f.write('# Exemplo de uso no HTML:\n')
        f.write('# <img src="{DOCUMENT_IMAGES[\'image1_png\']}" alt="Logo" />\n')
    
    print("✅ Arquivo criado: extracted_images_data.py")


def main():
    """Função principal"""
    print("="*70)
    print("  EXTRAÇÃO COMPLETA DO TEMPLATE WORD")
    print("  Cabeçalhos, Rodapés, Logos, Formas, Imagens")
    print("="*70)
    
    if not TEMPLATE_PATH.exists():
        print(f"\n❌ Arquivo não encontrado: {TEMPLATE_PATH}")
        return
    
    print(f"\n📂 Analisando: {TEMPLATE_PATH}\n")
    
    # 1. Cabeçalhos e rodapés
    extract_header_footer(TEMPLATE_PATH)
    
    # 2. Imagens (todas)
    images = extract_images_advanced(TEMPLATE_PATH)
    print(f"\n✅ Total de imagens extraídas: {len(images)}")
    
    # 3. Formas e desenhos
    extract_shapes_and_drawing(TEMPLATE_PATH)
    
    # 4. Salvar imagens
    if images:
        images = save_images_to_assets(images)
        generate_python_code_with_images(images)
    
    print("\n" + "="*70)
    print("🎉 EXTRAÇÃO COMPLETA!")
    print("="*70)
    print("\nArquivos gerados:")
    print("  1. assets/extracted_images/ - Imagens salvas")
    print("  2. extracted_images_data.py - Código Python com base64")
    print("\nAgora você pode:")
    print("  - Ver as imagens extraídas na pasta assets/")
    print("  - Usar o código Python gerado no template HTML")


if __name__ == '__main__':
    try:
        main()
    except Exception as e:
        print(f"\n❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
