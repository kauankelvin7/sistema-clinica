"""
Script para analisar e extrair estrutura do documento Word template
Extrai: imagens, estilos, formatação, estrutura completa
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor
import base64
from io import BytesIO

# Caminho do template
TEMPLATE_PATH = Path(__file__).parent / 'models' / 'modelo homologação.docx'

def extract_images_from_document(doc_path):
    """Extrai todas as imagens do documento Word"""
    doc = Document(doc_path)
    images = []
    
    print("📸 Extraindo imagens do documento...\n")
    
    # Iterar sobre relationships para encontrar imagens
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            
            # Converter para base64
            base64_image = base64.b64encode(image_data).decode('utf-8')
            
            # Detectar tipo MIME
            ext = rel.target_ref.split('.')[-1].lower()
            mime_types = {
                'png': 'image/png',
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'gif': 'image/gif',
                'bmp': 'image/bmp',
                'emf': 'image/x-emf',
                'wmf': 'image/x-wmf',
            }
            mime = mime_types.get(ext, 'image/png')
            
            images.append({
                'name': rel.target_ref,
                'mime': mime,
                'base64': f"data:{mime};base64,{base64_image}",
                'size': len(image_data)
            })
            
            print(f"✅ Imagem encontrada: {rel.target_ref}")
            print(f"   Tipo: {mime}")
            print(f"   Tamanho: {len(image_data):,} bytes")
            print()
    
    return images


def analyze_document_structure(doc_path):
    """Analisa estrutura completa do documento"""
    doc = Document(doc_path)
    
    print("📋 ANÁLISE COMPLETA DO DOCUMENTO\n")
    print("="*70)
    
    # Analisar parágrafos
    print("\n📝 PARÁGRAFOS E ESTILOS:\n")
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            print(f"Parágrafo {i}:")
            print(f"  Texto: {para.text[:80]}{'...' if len(para.text) > 80 else ''}")
            print(f"  Estilo: {para.style.name}")
            print(f"  Alinhamento: {para.alignment}")
            
            # Analisar runs (formatação inline)
            if para.runs:
                run = para.runs[0]
                print(f"  Fonte: {run.font.name if run.font.name else 'Padrão'}")
                print(f"  Tamanho: {run.font.size.pt if run.font.size else 'Padrão'} pt")
                print(f"  Negrito: {run.font.bold}")
                print(f"  Itálico: {run.font.italic}")
            print()
    
    # Analisar tabelas
    print("\n📊 TABELAS:\n")
    for i, table in enumerate(doc.tables, 1):
        print(f"Tabela {i}:")
        print(f"  Linhas: {len(table.rows)}")
        print(f"  Colunas: {len(table.columns)}")
        print(f"  Estilo: {table.style.name if table.style else 'Sem estilo'}")
        
        # Mostrar conteúdo da primeira linha
        if table.rows:
            print(f"  Primeira linha:")
            for j, cell in enumerate(table.rows[0].cells):
                print(f"    Célula {j}: {cell.text[:50]}")
        print()
    
    # Analisar seções
    print("\n📄 CONFIGURAÇÃO DE PÁGINA:\n")
    for i, section in enumerate(doc.sections, 1):
        print(f"Seção {i}:")
        print(f"  Largura: {section.page_width.inches:.2f} polegadas")
        print(f"  Altura: {section.page_height.inches:.2f} polegadas")
        print(f"  Margem Superior: {section.top_margin.inches:.2f} polegadas")
        print(f"  Margem Inferior: {section.bottom_margin.inches:.2f} polegadas")
        print(f"  Margem Esquerda: {section.left_margin.inches:.2f} polegadas")
        print(f"  Margem Direita: {section.right_margin.inches:.2f} polegadas")
        print()


def export_full_structure(doc_path, output_file='template_structure.txt'):
    """Exporta estrutura completa para arquivo"""
    doc = Document(doc_path)
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("="*70 + "\n")
        f.write("ESTRUTURA COMPLETA DO TEMPLATE\n")
        f.write("="*70 + "\n\n")
        
        # Exportar todo o texto
        f.write("CONTEÚDO COMPLETO:\n")
        f.write("-"*70 + "\n\n")
        
        for i, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():
                f.write(f"[Parágrafo {i} - Estilo: {para.style.name}]\n")
                f.write(f"{para.text}\n\n")
        
        # Exportar tabelas
        f.write("\n" + "="*70 + "\n")
        f.write("TABELAS:\n")
        f.write("="*70 + "\n\n")
        
        for i, table in enumerate(doc.tables, 1):
            f.write(f"TABELA {i}:\n")
            f.write("-"*70 + "\n")
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    f.write(f"[{row_idx},{col_idx}] {cell.text}\n")
                f.write("\n")
            f.write("\n")
    
    print(f"✅ Estrutura exportada para: {output_file}")


def save_images_to_assets(images):
    """Salva imagens extraídas na pasta assets"""
    assets_dir = Path('assets') / 'template_images'
    assets_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"\n💾 Salvando imagens em: {assets_dir}\n")
    
    for i, img in enumerate(images, 1):
        # Detectar extensão
        ext = img['name'].split('.')[-1]
        filename = f"logo_{i}.{ext}"
        filepath = assets_dir / filename
        
        # Decodificar base64 e salvar
        import re
        base64_data = re.sub('^data:image/.+;base64,', '', img['base64'])
        image_data = base64.b64decode(base64_data)
        
        with open(filepath, 'wb') as f:
            f.write(image_data)
        
        print(f"✅ {filename} salvo ({img['size']:,} bytes)")
        img['saved_path'] = str(filepath)
    
    return images


def main():
    """Função principal"""
    print("="*70)
    print("  ANALISADOR DE TEMPLATE WORD")
    print("  Extraindo estrutura, estilos e imagens")
    print("="*70 + "\n")
    
    if not TEMPLATE_PATH.exists():
        print(f"❌ Arquivo não encontrado: {TEMPLATE_PATH}")
        return
    
    print(f"📂 Analisando: {TEMPLATE_PATH}\n")
    
    # 1. Extrair imagens
    images = extract_images_from_document(TEMPLATE_PATH)
    print(f"\n✅ Total de imagens encontradas: {len(images)}\n")
    
    # 2. Salvar imagens
    if images:
        images = save_images_to_assets(images)
    
    # 3. Analisar estrutura
    analyze_document_structure(TEMPLATE_PATH)
    
    # 4. Exportar estrutura completa
    export_full_structure(TEMPLATE_PATH, 'template_structure.txt')
    
    # 5. Gerar código Python com dados extraídos
    print("\n" + "="*70)
    print("📝 Gerando código Python com imagens base64...")
    print("="*70 + "\n")
    
    with open('template_images_base64.py', 'w', encoding='utf-8') as f:
        f.write('"""\nImagens extraídas do template Word\n"""\n\n')
        f.write('TEMPLATE_IMAGES = {\n')
        for i, img in enumerate(images, 1):
            f.write(f'    "logo_{i}": "{img["base64"]}",\n')
        f.write('}\n')
    
    print("✅ Arquivo criado: template_images_base64.py")
    print("\n🎉 Análise completa! Agora você pode:")
    print("   1. Ver a estrutura em: template_structure.txt")
    print("   2. Ver as imagens em: assets/template_images/")
    print("   3. Usar as imagens base64 em: template_images_base64.py")


if __name__ == '__main__':
    try:
        main()
    except Exception as e:
        print(f"\n❌ Erro: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
