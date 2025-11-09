"""
Módulo de conversão HTML para DOCX
Sistema de Homologação de Atestados Médicos
Autor: Kauan Kelvin
Data: 09/11/2025

Este módulo implementa:
- Conversão de HTML para formato Word (.docx)
- Preservação de formatação e estrutura
- Suporte para imagens base64
"""

import os
import logging
import re
import base64
from pathlib import Path
from typing import Optional
from io import BytesIO

logger = logging.getLogger(__name__)


class DOCXConversionError(Exception):
    """Exceção para erros de conversão DOCX"""
    pass


def convert_html_to_docx_htmldocx(html_content: str, output_path: str) -> str:
    """
    Converte HTML para DOCX usando htmldocx
    
    Args:
        html_content: Conteúdo HTML
        output_path: Caminho do DOCX de saída
        
    Returns:
        str: Caminho do DOCX gerado
    """
    try:
        from htmldocx import HtmlToDocx
        from docx import Document
        
        logger.info("Convertendo HTML para DOCX com htmldocx...")
        
        # Garantir diretório de saída
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Criar documento vazio
        document = Document()
        
        # Converter HTML para DOCX
        parser = HtmlToDocx()
        parser.add_html_to_document(html_content, document)
        
        # Salvar
        document.save(output_path)
        
        if not os.path.exists(output_path):
            raise DOCXConversionError("DOCX não foi gerado")
        
        logger.info(f"✅ DOCX gerado com htmldocx: {output_path}")
        return output_path
        
    except ImportError:
        raise DOCXConversionError(
            "htmldocx não instalado. Instale com: pip install htmldocx"
        )
    except Exception as e:
        logger.error(f"Erro com htmldocx: {e}")
        raise DOCXConversionError(f"Erro htmldocx: {e}")


def convert_html_to_docx_pypandoc(html_content: str, output_path: str) -> str:
    """
    Converte HTML para DOCX usando pypandoc (requer pandoc instalado)
    
    Args:
        html_content: Conteúdo HTML
        output_path: Caminho do DOCX de saída
        
    Returns:
        str: Caminho do DOCX gerado
    """
    try:
        import pypandoc
        
        logger.info("Convertendo HTML para DOCX com pypandoc...")
        
        # Garantir diretório de saída
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Converter
        pypandoc.convert_text(
            html_content,
            'docx',
            format='html',
            outputfile=output_path,
            extra_args=['--standalone']
        )
        
        if not os.path.exists(output_path):
            raise DOCXConversionError("DOCX não foi gerado")
        
        logger.info(f"✅ DOCX gerado com pypandoc: {output_path}")
        return output_path
        
    except ImportError:
        raise DOCXConversionError(
            "pypandoc não instalado. Instale com: pip install pypandoc\n"
            "Também instale pandoc: https://pandoc.org/installing.html"
        )
    except Exception as e:
        logger.error(f"Erro com pypandoc: {e}")
        raise DOCXConversionError(f"Erro pypandoc: {e}")


def convert_html_to_docx_mammoth(html_content: str, output_path: str) -> str:
    """
    Converte HTML para DOCX usando mammoth (conversão reversa)
    
    Nota: mammoth é principalmente para DOCX->HTML, mas pode fazer o inverso
    com limitações
    
    Args:
        html_content: Conteúdo HTML
        output_path: Caminho do DOCX de saída
        
    Returns:
        str: Caminho do DOCX gerado
    """
    try:
        # mammoth não suporta HTML->DOCX diretamente
        # Esta é uma implementação básica usando python-docx
        from docx import Document
        from docx.shared import Pt, Inches
        from bs4 import BeautifulSoup
        
        logger.info("Convertendo HTML para DOCX (conversão básica)...")
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Criar documento
        document = Document()
        
        # Processar título
        title = soup.find('div', class_='doc-title')
        if title:
            p = document.add_paragraph(title.get_text(strip=True))
            p.style = 'Heading 1'
        
        # Processar conteúdo
        content = soup.find('div', class_='content')
        if content:
            for p_tag in content.find_all('p'):
                text = p_tag.get_text(strip=True)
                if text:
                    document.add_paragraph(text)
        
        # Processar tabela
        table = soup.find('table', class_='info-table')
        if table:
            rows = table.find_all('tr')
            if rows:
                doc_table = document.add_table(rows=len(rows), cols=2)
                doc_table.style = 'Light Grid Accent 1'
                
                for i, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    if len(cells) >= 2:
                        doc_table.rows[i].cells[0].text = cells[0].get_text(strip=True)
                        doc_table.rows[i].cells[1].text = cells[1].get_text(strip=True)
        
        # Salvar
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        
        if not os.path.exists(output_path):
            raise DOCXConversionError("DOCX não foi gerado")
        
        logger.info(f"✅ DOCX gerado (conversão básica): {output_path}")
        return output_path
        
    except ImportError:
        raise DOCXConversionError(
            "Dependências não instaladas. Instale com: pip install python-docx beautifulsoup4"
        )
    except Exception as e:
        logger.error(f"Erro na conversão básica: {e}")
        raise DOCXConversionError(f"Erro conversão básica: {e}")


def convert_html_to_docx(html_content: str, output_path: str,
                        method: Optional[str] = None) -> str:
    """
    Converte HTML para DOCX usando a melhor estratégia disponível
    
    Ordem de prioridade (se method não especificado):
    1. htmldocx (melhor preservação de formatação)
    2. pypandoc (boa qualidade, requer pandoc)
    3. Conversão básica (fallback)
    
    Args:
        html_content: Conteúdo HTML
        output_path: Caminho do DOCX de saída
        method: Método específico ('htmldocx', 'pypandoc', 'basic') ou None para auto
        
    Returns:
        str: Caminho do DOCX gerado
        
    Raises:
        DOCXConversionError: Se nenhum método funcionar
    """
    errors = []
    
    # Se método específico foi solicitado
    if method:
        method = method.lower()
        if method == 'htmldocx':
            return convert_html_to_docx_htmldocx(html_content, output_path)
        elif method == 'pypandoc':
            return convert_html_to_docx_pypandoc(html_content, output_path)
        elif method == 'basic':
            return convert_html_to_docx_mammoth(html_content, output_path)
        else:
            raise DOCXConversionError(f"Método inválido: {method}")
    
    # Tentar métodos na ordem de prioridade
    methods = [
        ('htmldocx', convert_html_to_docx_htmldocx),
        ('pypandoc', convert_html_to_docx_pypandoc),
        ('basic', convert_html_to_docx_mammoth),
    ]
    
    for name, converter in methods:
        try:
            logger.info(f"Tentando conversão com {name}...")
            return converter(html_content, output_path)
        except DOCXConversionError as e:
            error_msg = f"{name}: {e}"
            logger.warning(error_msg)
            errors.append(error_msg)
    
    # Se todos falharam
    error_details = "\n".join(errors)
    raise DOCXConversionError(
        f"Nenhum conversor DOCX disponível. Instale pelo menos um:\n"
        f"• htmldocx (recomendado): pip install htmldocx\n"
        f"• pypandoc: pip install pypandoc + pandoc\n"
        f"• basic: pip install python-docx beautifulsoup4\n\n"
        f"Erros:\n{error_details}"
    )


def convert_html_file_to_docx(html_path: str, output_path: Optional[str] = None,
                              method: Optional[str] = None) -> str:
    """
    Converte arquivo HTML para DOCX
    
    Args:
        html_path: Caminho do arquivo HTML
        output_path: Caminho do DOCX de saída (opcional, usa mesmo nome)
        method: Método de conversão (opcional)
        
    Returns:
        str: Caminho do DOCX gerado
    """
    # Ler HTML
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Determinar caminho de saída
    if not output_path:
        output_path = str(Path(html_path).with_suffix('.docx'))
    
    # Converter
    return convert_html_to_docx(html_content, output_path, method)


if __name__ == '__main__':
    # Teste rápido
    import sys
    
    if len(sys.argv) < 2:
        print("Uso: python html_to_docx.py <arquivo.html> [saida.docx]")
        sys.exit(1)
    
    html_file = sys.argv[1]
    docx_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        result = convert_html_file_to_docx(html_file, docx_file)
        print(f"✅ DOCX gerado: {result}")
    except DOCXConversionError as e:
        print(f"❌ Erro: {e}")
        sys.exit(1)
