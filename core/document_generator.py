"""
Módulo de geração de documentos Word (.docx)
Sistema de Homologação de Atestados Médicos
Autor: Kauan Kelvin
Data: 28/10/2025

Este módulo implementa:
- Geração segura de documentos a partir de templates
- Sanitização de nomes de arquivos
- Abertura automática de documentos gerados
- Tratamento robusto de erros
- Logging de operações
"""

from docx import Document
import os
import re
import logging
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional, Any

# Importar configurações centralizadas
try:
    from .config import (
        TEMPLATE_FILE, GENERATED_DOCS_DIR, DOCUMENT_PREFIX,
        DOCUMENT_EXTENSION, MAX_FILENAME_LENGTH
    )
except ImportError:
    # Fallback se config não estiver disponível
    TEMPLATE_FILE = Path(os.path.dirname(os.path.dirname(__file__))) / 'models' / 'modelo homologação.docx'
    GENERATED_DOCS_DIR = Path(os.path.dirname(os.path.dirname(__file__))) / 'data' / 'generated_documents'
    DOCUMENT_PREFIX = "Declaracao_"
    DOCUMENT_EXTENSION = ".docx"
    MAX_FILENAME_LENGTH = 200

# Configurar logger
logger = logging.getLogger(__name__)

class DocumentGenerationError(Exception):
    """Exceção personalizada para erros de geração de documentos"""
    pass

def sanitizar_nome_arquivo(nome: str, max_length: int = MAX_FILENAME_LENGTH) -> str:
    """
    Sanitiza um nome de arquivo removendo caracteres inválidos e limitando o tamanho.
    
    Args:
        nome: Nome original do arquivo
        max_length: Tamanho máximo permitido para o nome
        
    Returns:
        str: Nome de arquivo sanitizado e seguro
    """
    if not nome:
        return "Arquivo_Sem_Nome"
    
    # Remove caracteres inválidos para nomes de arquivo no Windows e Linux
    nome_limpo = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '', nome)
    
    # Remove espaços extras e substitui espaços por underscores
    nome_limpo = '_'.join(nome_limpo.split())
    
    # Remove pontos no início e fim (problemático no Windows)
    nome_limpo = nome_limpo.strip('.')
    
    # Limita o tamanho
    if len(nome_limpo) > max_length:
        nome_limpo = nome_limpo[:max_length]
    
    # Garante que não fique vazio após sanitização
    if not nome_limpo:
        nome_limpo = "Arquivo_Sanitizado"
    
    logger.debug(f"Nome de arquivo sanitizado: '{nome}' -> '{nome_limpo}'")
    return nome_limpo

def validar_dados_documento(data: Dict[str, Any]) -> bool:
    """
    Valida se todos os dados necessários para geração do documento estão presentes.
    
    Args:
        data: Dicionário com os dados do documento
        
    Returns:
        bool: True se os dados são válidos, False caso contrário
    """
    campos_obrigatorios = [
        "nome_paciente", "tipo_doc_paciente", "numero_doc_paciente",
        "data_atestado", "qtd_dias_atestado", "codigo_cid",
        "nome_medico", "tipo_registro_medico", "crm__medico", "uf_crm_medico"
    ]
    
    for campo in campos_obrigatorios:
        if campo not in data or not str(data[campo]).strip():
            logger.warning(f"Campo obrigatório ausente ou vazio: {campo}")
            return False
    
    # Validar que qtd_dias_atestado é numérico
    try:
        dias = int(data["qtd_dias_atestado"])
        if dias <= 0:
            logger.warning("Quantidade de dias deve ser maior que zero")
            return False
    except (ValueError, TypeError):
        logger.warning("Quantidade de dias deve ser um número inteiro")
        return False
    
    return True

def substituir_placeholders(document: Document, replacements: Dict[str, str]) -> None:
    """
    Substitui placeholders no documento Word por valores reais.
    Itera por parágrafos e tabelas.
    
    IMPORTANTE: Placeholders podem estar divididos em múltiplos runs.
    Esta função reconstrói o texto completo do parágrafo antes de substituir.
    
    Args:
        document: Objeto Document do python-docx
        replacements: Dicionário com os placeholders e seus valores
    """
    def substituir_em_paragrafo(paragraph):
        """Substitui placeholders em um parágrafo preservando formatação"""
        # Obter texto completo do parágrafo
        texto_completo = paragraph.text
        
        # Verificar se há placeholders para substituir
        precisa_substituir = False
        for key in replacements.keys():
            if key in texto_completo:
                precisa_substituir = True
                break
        
        if not precisa_substituir:
            return
        
        # Fazer substituições no texto completo
        texto_novo = texto_completo
        # Substituir chaves maiores primeiro para evitar conflitos (ex: {nome}{crm} vs {nome})
        for key, value in sorted(replacements.items(), key=lambda kv: -len(kv[0])):
            texto_novo = texto_novo.replace(key, value)
        
        # Se o texto mudou, atualizar o parágrafo
        if texto_novo != texto_completo:
            # Limpar runs existentes
            for run in paragraph.runs:
                run.text = ""
            
            # Adicionar texto novo no primeiro run (preserva formatação base)
            if paragraph.runs:
                paragraph.runs[0].text = texto_novo
            else:
                paragraph.add_run(texto_novo)
    
    # Substituir em parágrafos principais
    for paragraph in document.paragraphs:
        substituir_em_paragrafo(paragraph)
    
    # Substituir em tabelas
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    substituir_em_paragrafo(paragraph)

def abrir_arquivo_automaticamente(caminho_arquivo: Path) -> bool:
    """
    Abre o arquivo gerado automaticamente com o aplicativo padrão do sistema.
    SOMENTE em ambiente de desenvolvimento local (não em produção/servidor).
    
    Args:
        caminho_arquivo: Caminho do arquivo a ser aberto
        
    Returns:
        bool: True se aberto com sucesso, False caso contrário
    """
    # Detectar se está rodando em servidor (Render, Vercel, etc)
    is_production = os.getenv('RENDER') or os.getenv('VERCEL') or os.getenv('RAILWAY')
    
    if is_production:
        logger.debug("Ambiente de produção detectado - não abrindo arquivo automaticamente")
        return False
    
    try:
        if os.name == 'nt':  # Windows
            os.startfile(str(caminho_arquivo))
        elif os.uname().sysname == 'Darwin':  # macOS
            subprocess.Popen(['open', str(caminho_arquivo)])
        else:  # Linux e outros Unix
            subprocess.Popen(['xdg-open', str(caminho_arquivo)])
        
        logger.info(f"Arquivo aberto automaticamente: {caminho_arquivo}")
        return True
        
    except Exception as e:
        logger.warning(f"Não foi possível abrir o arquivo automaticamente: {e}")
        return False

def generate_document(data: Dict[str, Any]) -> Optional[str]:
    """
    Carrega o modelo .docx, substitui os placeholders pelos dados fornecidos
    e salva o novo documento.
    
    Args:
        data: Dicionário contendo todos os dados necessários para o documento
        
    Returns:
        str: Caminho do arquivo gerado, ou None em caso de erro
        
    Raises:
        DocumentGenerationError: Se houver erro na geração do documento
    """
    try:
        logger.info("Iniciando geração de documento...")
        
        # Validar dados de entrada
        if not validar_dados_documento(data):
            raise DocumentGenerationError("Dados de entrada inválidos ou incompletos")
        
        # Verificar se o template existe
        if not TEMPLATE_FILE.exists():
            raise DocumentGenerationError(f"Arquivo de template não encontrado: {TEMPLATE_FILE}")
        
        # Garantir que a pasta de saída existe
        GENERATED_DOCS_DIR.mkdir(parents=True, exist_ok=True)
        
        # Carregar o documento template
        logger.debug(f"Carregando template: {TEMPLATE_FILE}")
        document = Document(str(TEMPLATE_FILE))
        
        # Preparar substituições (cada placeholder separadamente)
        nome_med = str(data.get("nome_medico", "")).strip()
        tipo_reg = str(data.get('tipo_registro_medico', '')).strip()
        crm_num = str(data.get('crm__medico', '')).strip()
        uf_crm = str(data.get('uf_crm_medico', '')).strip()

        # Normalizar nome do médico: remover títulos para evitar duplicação (Dr., Dra., Dr(a), Prof., etc.)
        # Mantemos somente o nome e sobrenome; o template pode conter o prefixo 'Dr. (a) '
        nome_sem_titulo = re.sub(r'(?i)^\s*(dr\.?|dra\.?|dr\(a\)\.?|dr\(a\)|prof\.?|profa\.?|sr\.?|sra\.?|med\.?\s*)\s*', '', nome_med)
        nome_sem_titulo = nome_sem_titulo.strip()

        # Substituição combinada para o template que usa {nome_medico}{crm__medico}-{uf_crm_medico}
        # Formatamos sem repetir o prefixo 'Dr.' caso o template já forneça esse texto.
        # Exemplo final esperado dentro do template: 'Dr. (a) Nome Sobrenome CRM 12345-DF'
        partes_registro = []
        if tipo_reg:
            partes_registro.append(tipo_reg)
        if crm_num:
            partes_registro.append(crm_num)
        registro_formatado_sem_uf = ' '.join(partes_registro)  # Ex: "CRM 12345"
        
        # Para o placeholder combinado, incluir o UF
        registro_formatado_completo = registro_formatado_sem_uf
        if uf_crm and registro_formatado_sem_uf:
            registro_formatado_completo = f"{registro_formatado_sem_uf}-{uf_crm}"

        medico_completo = f"{nome_sem_titulo} {registro_formatado_completo}".strip()

        replacements = {
            "{nome_paciente}": str(data.get("nome_paciente", "")).strip(),
            "{documento_paciente_formatado}": f"{data.get('tipo_doc_paciente', '').upper()} nº: {data.get('numero_doc_paciente', '')}",
            "{data_atestado}": str(data.get("data_atestado", "")).strip(),
            "{qtd_dias_atestado}": str(data.get("qtd_dias_atestado", "")),
            "{código_cid}": str(data.get("codigo_cid", "")).strip(),
            "{cargo_paciente}": str(data.get("cargo_paciente", "")).strip(),
            "{empresa_paciente}": str(data.get("empresa_paciente", "")).strip(),
            "___/___/____": datetime.now().strftime("%d/%m/%Y"),
            # Para o placeholder combinado na tabela, usar nome normalizado (sem Dr.)
            "{nome_medico}{crm__medico}-{uf_crm_medico}": medico_completo,
            # Para placeholders individuais no parágrafo, usar valores originais/formatados
            "{nome_medico}": nome_sem_titulo,
            "{crm__medico}": registro_formatado_sem_uf if registro_formatado_sem_uf else crm_num,
            "{tipo_registro_medico}": tipo_reg,
            "{uf_crm_medico}": uf_crm
        }
        
        # Substituir placeholders
        logger.debug("Substituindo placeholders no documento")
        substituir_placeholders(document, replacements)
        
        # Gerar nome do arquivo de saída
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nome_paciente_sanitizado = sanitizar_nome_arquivo(
            data.get('nome_paciente', 'Paciente'),
            max_length=50  # Limitar nome do paciente no arquivo
        )
        output_filename = f"{DOCUMENT_PREFIX}{nome_paciente_sanitizado}_{timestamp}{DOCUMENT_EXTENSION}"
        output_path = GENERATED_DOCS_DIR / output_filename
        
        # Salvar documento
        logger.debug(f"Salvando documento em: {output_path}")
        document.save(str(output_path))
        logger.info(f"Documento gerado com sucesso: {output_path}")
        
        # Tentar abrir o arquivo automaticamente
        abrir_arquivo_automaticamente(output_path)
        
        return str(output_path)
        
    except Exception as e:
        logger.error(f"Erro ao gerar documento: {e}", exc_info=True)
        raise DocumentGenerationError(f"Erro ao gerar documento: {e}")

# Manter compatibilidade com código legado
def gerar_documento(*args, **kwargs):
    """Alias para generate_document (compatibilidade)"""
    return generate_document(*args, **kwargs)