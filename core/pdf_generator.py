"""
Módulo de conversão DOCX para PDF
Sistema de Homologação de Atestados Médicos
Autor: Kauan Kelvin
Data: 30/10/2025

Este módulo converte documentos DOCX para PDF usando docx2pdf (Windows/COM)
ou reportlab como fallback multiplataforma.
"""

import os
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

class PDFGenerationError(Exception):
    """Exceção personalizada para erros de geração de PDF"""
    pass


def convert_docx_to_pdf_com(docx_path: str) -> str:
    """
    Converte DOCX para PDF usando Word COM automation (Windows apenas).
    Preserva TODA formatação original.
    
    Args:
        docx_path: Caminho do arquivo DOCX
        
    Returns:
        str: Caminho do arquivo PDF gerado
    """
    try:
        from docx2pdf import convert
        
        # Caminho do PDF (mesmo nome, extensão .pdf)
        pdf_path = str(Path(docx_path).with_suffix('.pdf'))
        
        # Converter usando Word COM
        convert(docx_path, pdf_path)
        
        logger.info(f"PDF gerado via Word COM: {pdf_path}")
        return pdf_path
        
    except Exception as e:
        logger.warning(f"Erro ao usar docx2pdf: {e}")
        raise


def convert_docx_to_pdf_reportlab(docx_path: str) -> str:
    """
    Converte DOCX para PDF usando reportlab (multiplataforma).
    Extrai texto e formatação básica do DOCX.
    
    Args:
        docx_path: Caminho do arquivo DOCX
        
    Returns:
        str: Caminho do arquivo PDF gerado
    """
    try:
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
        from reportlab.lib import colors
        
        # Abrir DOCX
        doc = Document(docx_path)
        
        # Caminho do PDF
        pdf_path = str(Path(docx_path).with_suffix('.pdf'))
        
        # Criar PDF
        pdf_doc = SimpleDocTemplate(
            pdf_path,
            pagesize=A4,
            rightMargin=2.5*cm,
            leftMargin=2.5*cm,
            topMargin=2.5*cm,
            bottomMargin=2.5*cm
        )
        
        # Estilos
        styles = getSampleStyleSheet()
        
        # Criar estilos customizados baseados no Word
        style_title = ParagraphStyle(
            'WordTitle',
            parent=styles['Heading1'],
            fontSize=14,
            fontName='Helvetica-Bold',
            alignment=TA_CENTER,
            spaceAfter=20,
            spaceBefore=10
        )
        
        style_normal = ParagraphStyle(
            'WordNormal',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Helvetica',
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            leading=16
        )
        
        style_center = ParagraphStyle(
            'WordCenter',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Helvetica',
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        # Story (elementos do PDF)
        story = []
        
        # Processar parágrafos do DOCX
        for para in doc.paragraphs:
            if not para.text.strip():
                story.append(Spacer(1, 0.3*cm))
                continue
            
            # Detectar alinhamento
            alignment = para.alignment
            if alignment == 1:  # CENTER
                style = style_center
            elif alignment == 3:  # JUSTIFY
                style = style_normal
            else:
                style = style_normal
            
            # Detectar se é título (negrito ou tamanho maior)
            is_bold = any(run.bold for run in para.runs)
            if is_bold and len(para.text) < 50:
                style = style_title
            
            # Adicionar parágrafo
            story.append(Paragraph(para.text, style))
        
        # Processar tabelas do DOCX
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = '\n'.join([p.text for p in cell.paragraphs])
                    row_data.append(Paragraph(cell_text, style_center))
                table_data.append(row_data)
            
            # Criar tabela no PDF
            if table_data:
                pdf_table = Table(table_data, colWidths=[15*cm])
                pdf_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                    ('TOPPADDING', (0, 0), (-1, -1), 12),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
                ]))
                story.append(pdf_table)
                story.append(Spacer(1, 0.5*cm))
        
        # Construir PDF
        pdf_doc.build(story)
        
        logger.info(f"PDF gerado via reportlab: {pdf_path}")
        return pdf_path
        
    except Exception as e:
        logger.error(f"Erro ao converter com reportlab: {e}")
        raise


def generate_pdf_direct(data: dict, output_dir: Optional[str] = None) -> str:
    """
    Gera PDF a partir de DOCX (mantém formatação original).
    
    Fluxo:
    1. Gera DOCX usando document_generator
    2. Converte DOCX → PDF (preservando formatação)
    
    Args:
        data: Dicionário com os dados do documento
        output_dir: Diretório de saída
        
    Returns:
        str: Caminho do arquivo PDF gerado
    """
    try:
        from core.document_generator import generate_document
        
        # 1. Gerar DOCX primeiro
        docx_path = generate_document(data)
        if not docx_path or not os.path.exists(docx_path):
            raise PDFGenerationError("Falha ao gerar documento Word")
        
        logger.info(f"DOCX gerado: {docx_path}")
        
        # 2. Converter para PDF
        # Tentar docx2pdf primeiro (Windows - preserva 100% formatação)
        try:
            pdf_path = convert_docx_to_pdf_com(docx_path)
            return pdf_path
        except Exception as e:
            logger.warning(f"docx2pdf falhou (tentando reportlab): {e}")
        
        # Fallback: reportlab (multiplataforma)
        try:
            pdf_path = convert_docx_to_pdf_reportlab(docx_path)
            return pdf_path
        except Exception as e:
            logger.error(f"Todas conversões falharam: {e}")
            raise PDFGenerationError(f"Não foi possível converter para PDF: {e}")
        
    except PDFGenerationError:
        raise
    except Exception as e:
        logger.error(f"Erro ao gerar PDF: {e}", exc_info=True)
        raise PDFGenerationError(f"Erro ao gerar PDF: {e}")
