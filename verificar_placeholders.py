"""Script para verificar placeholders no modelo Word"""
from docx import Document

doc = Document('models/modelo homologação.docx')

print("=" * 80)
print("PLACEHOLDERS ENCONTRADOS NO MODELO")
print("=" * 80)

print("\n📄 PARÁGRAFOS:")
print("-" * 80)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"{i}: {p.text}")

print("\n📊 TABELAS:")
print("-" * 80)
for i, table in enumerate(doc.tables):
    print(f"\nTabela {i}:")
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f"  Linha {row_idx}, Célula {cell_idx}: {cell.text}")

print("\n" + "=" * 80)
