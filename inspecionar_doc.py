from docx import Document
p='data/generated_documents/Declaracao_TESTE_JOÃO_SILVA_20251030_080448.docx'
d=Document(p)
# find paragraph that contains 'IDENTIFICAÇÃO' or 'Dr.'
for i,table in enumerate(d.tables):
    for r,row in enumerate(table.rows):
        for c,cell in enumerate(row.cells):
            if 'IDENTIFICAÇÃO' in cell.text or 'Dr.' in cell.text:
                print('Table',i,'Row',r,'Cell',c)
                for pi,para in enumerate(cell.paragraphs):
                    print(' PAR',pi,repr(para.text))
                    for ri,run in enumerate(para.runs):
                        print('  RUN',ri,repr(run.text))
