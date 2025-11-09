#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Extrair Fontes do Documento Word
=================================
Analisa todas as fontes e tamanhos usados no template Word
"""

from docx import Document
from docx.shared import Pt

def extract_fonts():
    """Extrai fontes e tamanhos do Word"""
    doc = Document('models/modelo homologação.docx')
    
    fonts_data = {}
    
    print("\n" + "="*60)
    print("  ANÁLISE DE FONTES DO DOCUMENTO WORD")
    print("="*60)
    print()
    
    # Analisar parágrafos
    for para in doc.paragraphs:
        if para.runs:
            for run in para.runs:
                font_name = run.font.name or 'Calibri'
                font_size = run.font.size.pt if run.font.size else 11
                text_sample = run.text.strip()[:50]
                
                if text_sample:
                    key = (font_name, font_size)
                    if key not in fonts_data:
                        fonts_data[key] = []
                    fonts_data[key].append(text_sample)
    
    # Analisar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.runs:
                        for run in para.runs:
                            font_name = run.font.name or 'Calibri'
                            font_size = run.font.size.pt if run.font.size else 11
                            text_sample = run.text.strip()[:50]
                            
                            if text_sample:
                                key = (font_name, font_size)
                                if key not in fonts_data:
                                    fonts_data[key] = []
                                fonts_data[key].append(text_sample)
    
    # Mostrar resultados
    print("📝 FONTES ENCONTRADAS:\n")
    for (font_name, font_size), samples in sorted(fonts_data.items()):
        print(f"✓ {font_name} - {font_size}pt")
        print(f"  Exemplos: {samples[0][:60]}")
        if len(samples) > 1:
            print(f"           {samples[1][:60]}")
        print()
    
    # Resumo
    print("="*60)
    print("📊 RESUMO:")
    print(f"   Total de combinações fonte/tamanho: {len(fonts_data)}")
    print()
    
    # Fontes mais comuns
    font_usage = {}
    for (font_name, font_size) in fonts_data.keys():
        if font_name not in font_usage:
            font_usage[font_name] = []
        font_usage[font_name].append(font_size)
    
    print("🔤 FONTES PRINCIPAIS:")
    for font_name, sizes in sorted(font_usage.items()):
        sizes_str = ', '.join([f"{s}pt" for s in sorted(set(sizes))])
        print(f"   • {font_name}: {sizes_str}")
    print()

if __name__ == "__main__":
    extract_fonts()
