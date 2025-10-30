"""Script de teste para verificar geração de documento"""
from datetime import datetime
from core.document_generator import generate_document

# Dados de teste
data = {
    "nome_paciente": "TESTE JOÃO SILVA",
    "tipo_doc_paciente": "CPF",
    "numero_doc_paciente": "123.456.789-00",
    "cargo_paciente": "Desenvolvedor",
    "empresa_paciente": "Tech Corp LTDA",
    "data_atestado": "25/10/2025",
    "qtd_dias_atestado": 3,
    "codigo_cid": "Z76.5",
    "nome_medico": "Dr. Maria Santos",
    "tipo_registro_medico": "CRM",
    "crm__medico": "12345",
    "uf_crm_medico": "DF"
}

print("=" * 80)
print("TESTE DE GERAÇÃO DE DOCUMENTO")
print("=" * 80)
print("\n📋 Dados de entrada:")
for key, value in data.items():
    print(f"  {key}: {value}")

print("\n🔄 Gerando documento...")
try:
    output_path = generate_document(data)
    print(f"\n✅ Documento gerado com sucesso!")
    print(f"📁 Caminho: {output_path}")
    
    # Verificar se foi preenchido
    from docx import Document
    doc = Document(output_path)
    
    print("\n📄 Conteúdo do documento gerado:")
    print("-" * 80)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"{i}: {p.text[:150]}")
    
    print("\n📊 Tabelas:")
    for i, table in enumerate(doc.tables):
        print(f"\nTabela {i}:")
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    print(f"  {cell.text[:100]}")
    
except Exception as e:
    print(f"\n❌ ERRO: {e}")
    import traceback
    traceback.print_exc()

print("\n" + "=" * 80)
